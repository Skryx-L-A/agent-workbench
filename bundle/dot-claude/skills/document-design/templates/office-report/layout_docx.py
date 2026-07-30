"""layout_docx.py -- the MECHANICS of a multi-page DOCX report: page setup, heading
styles, running head, footer page numbers, a table without a cage, font embedding.
This file is meant to be read and changed -- what is worth keeping across documents
is in here, what must be re-decided each time is in tokens.py.

Mirrors templates/report/layout.typ one level down: same split (mechanics here,
decisions in tokens.py, content in report.py), same gate, same table philosophy
(rules, not a grid -- see reference/antipatterns.md N1). Built with python-docx,
which has no notion of a "compiler" -- gate() is called explicitly at the top of
build(), because nothing else will stop an undecided document from being written.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

_here = Path(__file__).resolve().parent
for _candidate in (_here / "scripts", _here.parent / "scripts", _here.parent.parent / "scripts"):
    if (_candidate / "officefonts.py").exists():
        sys.path.insert(0, str(_candidate))
        break
else:
    sys.exit(
        "officefonts.py not found in ./scripts, ../scripts or ../../scripts next to this "
        "template.\nCopy scripts/officefonts.py alongside this template, or run it from "
        "inside the document-design project."
    )
import officefonts  # noqa: E402

GATE_MESSAGE = (
    "Layout for this document has not been decided yet.\n"
    "Open tokens.py, replace every value with your own decision for THIS document,\n"
    "write the reason on each `# why:` comment, then set DESIGN_DECIDED = True.\n"
    "Do not simply flip the flag: the shipped numbers are an example, not a house style."
)


def gate(tokens) -> None:
    if not getattr(tokens, "DESIGN_DECIDED", False):
        sys.exit(GATE_MESSAGE)
    for group in (tokens.FONT_TEXT, tokens.FONT_DISPLAY):
        if group["family"] == "CHOOSE-ME" or not Path(group["regular"]).exists():
            sys.exit(
                f"Font not decided or not found: {group}.\n"
                "Run scripts/make-static-fonts.sh on a variable font first, or point\n"
                "FONT_TEXT/FONT_DISPLAY at real Regular/Bold .ttf files."
            )


def _rgb(hex_str: str) -> RGBColor:
    return RGBColor.from_string(hex_str)


def _set_run_font(run, family: str, size_pt: float, color: str, bold=None, italic=None):
    run.font.name = family
    run.font.size = Pt(size_pt)
    run.font.color.rgb = _rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    # East-Asian / complex-script slots default back to Calibri/Times inside Word even
    # when the Latin slot is set correctly -- rPr/rFonts needs all four explicitly.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), family)


def _field_run(paragraph, field_code: str):
    """Insert a Word field (e.g. PAGE) as a run -- python-docx has no helper for this."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._element.append(fld_begin)
    run._element.append(instr)
    run._element.append(fld_end)
    return run


def _no_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def _hline(cell_row, tokens, weight_pt=0.75, color=None):
    """A horizontal rule under a row -- the table's only border, applied per cell
    because Word borders are a per-cell property, not a per-row one."""
    for cell in cell_row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(int(weight_pt * 8)))  # eighths of a point
        bottom.set(qn("w:color"), color or tokens.INK["rule"])
        borders.append(bottom)
        tc_pr.append(borders)


class H1:
    def __init__(self, text: str):
        self.text = text


class H2:
    def __init__(self, text: str):
        self.text = text


class H3:
    def __init__(self, text: str):
        self.text = text


class P:
    def __init__(self, text: str):
        self.text = text


class TABLE:
    def __init__(self, head: list[str], rows: list[list[str]]):
        self.head = head
        self.rows = rows


class FIGURE:
    def __init__(self, path: str, caption: str | None = None, width_mm: float | None = None):
        self.path = path
        self.caption = caption
        self.width_mm = width_mm


def _add_heading(doc, tokens, block):
    size, upper, level = {
        H1: (tokens.SIZE_PT["h1"], False, 1),
        H2: (tokens.SIZE_PT["h2"], False, 2),
        H3: (tokens.SIZE_PT["h3"], True, 3),
    }[type(block)]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True  # no orphan heading at a page foot
    text = block.text.upper() if upper else block.text
    run = p.add_run(text)
    _set_run_font(run, tokens.FONT_DISPLAY["family"], size, tokens.INK["text"], bold=True)
    if level == 3:
        run.font.size = Pt(size)


def _add_paragraph(doc, tokens, block):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = tokens.LINE_SPACING
    p.paragraph_format.space_after = Pt(tokens.SIZE_PT["body"] * (tokens.LINE_SPACING - 1) + 4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if tokens.JUSTIFY else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(block.text)
    _set_run_font(run, tokens.FONT_TEXT["family"], tokens.SIZE_PT["body"], tokens.INK["text"])


def _add_table(doc, tokens, block: TABLE):
    n_cols = len(block.head)
    table = doc.add_table(rows=1 + len(block.rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_borders(table)
    _hline(table.rows[0], tokens, weight_pt=1.0)  # rule under the head
    for j, h in enumerate(block.head):
        cell = table.rows[0].cells[j]
        run = cell.paragraphs[0].add_run(h)
        _set_run_font(run, tokens.FONT_DISPLAY["family"], tokens.SIZE_PT["small"], tokens.INK["text"], bold=True)
    for i, row in enumerate(block.rows, start=1):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            run = cell.paragraphs[0].add_run(str(val))
            _set_run_font(run, tokens.FONT_DISPLAY["family"], tokens.SIZE_PT["small"], tokens.INK["text"])
    _hline(table.rows[-1], tokens, weight_pt=1.0)  # rule under the last row -- no cage (N1)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _add_figure(doc, tokens, block: FIGURE):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    width = Mm(block.width_mm) if block.width_mm else None
    p.add_run().add_picture(block.path, width=width)
    if block.caption:
        cap = doc.add_paragraph()
        run = cap.add_run(block.caption)
        _set_run_font(run, tokens.FONT_DISPLAY["family"], tokens.SIZE_PT["small"], tokens.INK["soft"])


_DISPATCH = {H1: _add_heading, H2: _add_heading, H3: _add_heading, P: _add_paragraph, TABLE: _add_table, FIGURE: _add_figure}


def _accent_rule(doc, tokens, width_mm=20, weight_pt=3):
    """A drawn rule of an exact width, not a text glyph. A box-drawing character
    ("─") silently falls back to a system font for that one glyph if the chosen face
    doesn't contain it -- found on this template's own title page (Archivo has no
    U+2500; LibreOffice substituted Helvetica for just that run, invisible in the
    source, only caught by `pdffonts`/docrender). A 1-cell table with only a bottom
    border gives a vector line of a known width; a paragraph border runs edge to edge
    and cannot be shortened without an unrelated indent hack."""
    doc.add_paragraph().paragraph_format.space_before = Pt(96)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _no_borders(table)
    table.columns[0].width = Mm(width_mm)
    table.rows[0].cells[0].width = Mm(width_mm)
    _hline(table.rows[0], tokens, weight_pt=weight_pt, color=tokens.INK["accent"])
    table.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(0)


def _title_page(doc, tokens, title, subtitle, meta):
    _accent_rule(doc, tokens)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run(title)
    _set_run_font(run, tokens.FONT_DISPLAY["family"], tokens.SIZE_PT["title"], tokens.INK["text"], bold=True)

    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        _set_run_font(run, tokens.FONT_TEXT["family"], tokens.SIZE_PT["h2"], tokens.INK["soft"], italic=True)

    for _ in range(10):
        doc.add_paragraph()

    for key, val in meta.items():
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Mm(32))
        k = p.add_run(key.upper() + "\t")
        _set_run_font(k, tokens.FONT_DISPLAY["family"], tokens.SIZE_PT["small"], tokens.INK["soft"])
        v = p.add_run(str(val))
        _set_run_font(v, tokens.FONT_DISPLAY["family"], tokens.SIZE_PT["small"], tokens.INK["text"])

    doc.add_page_break()


def _running_head(section, tokens, title):
    section.different_first_page_header_footer = True
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(title)
    _set_run_font(run, tokens.FONT_DISPLAY["family"], tokens.SIZE_PT["small"], tokens.INK["soft"])

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    _set_run_font(run, tokens.FONT_DISPLAY["family"], tokens.SIZE_PT["small"], tokens.INK["soft"])
    _field_run(p, "PAGE")

    # First page carries neither -- title page already states what the document is.
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""


def build(tokens, blocks: list, meta: dict, out_path: str):
    """meta needs at least title; subtitle/author/date/version are shown on the
    title page. blocks is the declarative content list (H1/H2/H3/P/TABLE/FIGURE)."""
    gate(tokens)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(tokens.PAPER_MM[0])
    section.page_height = Mm(tokens.PAPER_MM[1])
    section.left_margin = Mm(tokens.MARGIN_MM["left"])
    section.right_margin = Mm(tokens.MARGIN_MM["right"])
    section.top_margin = Mm(tokens.MARGIN_MM["top"])
    section.bottom_margin = Mm(tokens.MARGIN_MM["bottom"])
    section.header_distance = Mm(tokens.HEADER_DISTANCE_MM)
    section.footer_distance = Mm(tokens.FOOTER_DISTANCE_MM)

    normal = doc.styles["Normal"]
    normal.font.name = tokens.FONT_TEXT["family"]
    normal.font.size = Pt(tokens.SIZE_PT["body"])
    normal.paragraph_format.line_spacing = tokens.LINE_SPACING

    _title_page(doc, tokens, meta.get("title", ""), meta.get("subtitle"), {
        k: v for k, v in meta.items() if k not in ("title", "subtitle")
    })
    _running_head(doc.sections[0], tokens, meta.get("title", ""))

    for block in blocks:
        _DISPATCH[type(block)](doc, tokens, block)

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)

    font_specs = []
    for group in (tokens.FONT_TEXT, tokens.FONT_DISPLAY):
        spec = {"family": group["family"], "regular": group["regular"]}
        if group.get("bold"):
            spec["bold"] = group["bold"]
        font_specs.append(spec)
    # Two families may share a family name only if identical; de-dupe by family.
    seen = {}
    for spec in font_specs:
        seen[spec["family"]] = spec
    officefonts.embed_fonts_docx(out_path, list(seen.values()))
